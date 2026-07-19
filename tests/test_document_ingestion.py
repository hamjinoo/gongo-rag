"""document_ingestion.py의 파일 형식별 추출 테스트."""

import sys
from io import BytesIO
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))
sys.stdout.reconfigure(encoding="utf-8")

from document_ingestion import (  # noqa: E402
    ExtractionConfig,
    InvalidDocumentError,
    OCRStatus,
    UnsupportedDocumentError,
    extract_document,
)


class FakeOCR:
    def __init__(self, text: str = "스캔 문서에서 읽은 글자") -> None:
        self.text = text
        self.calls = 0

    def status(self, language: str) -> OCRStatus:
        return OCRStatus(True, "fake", f"ready: {language}", ("kor", "eng"))

    def recognize(self, image_bytes: bytes, language: str, timeout_seconds: int) -> str:
        assert image_bytes
        assert language == "kor+eng"
        assert timeout_seconds > 0
        self.calls += 1
        return self.text


class UnavailableOCR:
    def status(self, language: str) -> OCRStatus:
        return OCRStatus(False, "fake", "OCR 엔진 없음")

    def recognize(self, image_bytes: bytes, language: str, timeout_seconds: int) -> str:
        from document_ingestion import OCRUnavailableError

        raise OCRUnavailableError("OCR 엔진 없음")


def make_native_pdf() -> bytes:
    import pymupdf

    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Native PDF text for document ingestion")
    data = document.tobytes()
    document.close()
    return data


def make_scanned_pdf() -> bytes:
    import pymupdf
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (800, 300), "white")
    draw = ImageDraw.Draw(image)
    draw.text((40, 120), "SCANNED DOCUMENT", fill="black")
    image_buffer = BytesIO()
    image.save(image_buffer, format="PNG")

    document = pymupdf.open()
    page = document.new_page(width=800, height=300)
    page.insert_image(page.rect, stream=image_buffer.getvalue())
    data = document.tobytes()
    document.close()
    return data


def make_locked_pdf() -> bytes:
    import pymupdf

    document = pymupdf.open()
    document.new_page().insert_text((72, 72), "locked")
    data = document.tobytes(
        encryption=pymupdf.PDF_ENCRYPT_AES_256,
        owner_pw="owner-password",
        user_pw="user-password",
    )
    document.close()
    return data


def make_docx() -> bytes:
    from docx import Document

    document = Document()
    document.add_paragraph("신청 자격은 창업 3년 이내 기업입니다.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "내용"
    table.cell(1, 0).text = "지원 금액"
    table.cell(1, 1).text = "최대 1억원"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_native_pdf_uses_embedded_text():
    fake_ocr = FakeOCR()
    result = extract_document(
        "notice.pdf",
        make_native_pdf(),
        config=ExtractionConfig(min_native_chars=10),
        ocr_backend=fake_ocr,
    )
    assert "Native PDF text" in result.text
    assert result.pages[0].method == "native"
    assert fake_ocr.calls == 0


def test_scanned_pdf_uses_ocr_fallback():
    fake_ocr = FakeOCR("스캔 PDF의 신청 기간은 7월 31일까지입니다.")
    result = extract_document(
        "scan.pdf",
        make_scanned_pdf(),
        config=ExtractionConfig(min_native_chars=10),
        ocr_backend=fake_ocr,
    )
    assert "7월 31일" in result.text
    assert result.pages[0].method == "ocr"
    assert result.used_ocr is True
    assert fake_ocr.calls == 1


def test_scanned_pdf_survives_missing_ocr_engine():
    result = extract_document(
        "scan.pdf",
        make_scanned_pdf(),
        ocr_backend=UnavailableOCR(),
    )
    assert result.text == ""
    assert result.pages[0].method == "empty"
    assert any("OCR 엔진 없음" in warning for warning in result.warnings)


def test_locked_pdf_has_clear_error():
    try:
        extract_document("locked.pdf", make_locked_pdf())
    except InvalidDocumentError as exc:
        assert "암호로 잠긴 PDF" in str(exc)
    else:
        raise AssertionError("암호로 잠긴 PDF를 거절해야 합니다.")


def test_docx_extracts_paragraphs_and_tables():
    result = extract_document("notice.docx", make_docx())
    assert "창업 3년 이내" in result.text
    assert "[표]" in result.text
    assert "지원 금액 | 최대 1억원" in result.text
    assert result.pages[0].method == "docx"


def test_direct_image_uses_same_ocr_backend():
    from PIL import Image

    buffer = BytesIO()
    Image.new("RGB", (20, 20), "white").save(buffer, format="PNG")
    fake_ocr = FakeOCR("이미지에서 추출한 글자")
    result = extract_document("photo.png", buffer.getvalue(), ocr_backend=fake_ocr)
    assert result.text == "이미지에서 추출한 글자"
    assert result.used_ocr is True
    assert fake_ocr.calls == 1


def test_broken_image_has_clear_error():
    try:
        extract_document("broken.png", b"not a real image", ocr_backend=FakeOCR())
    except InvalidDocumentError as exc:
        assert "이미지 파일을 읽을 수 없습니다" in str(exc)
    else:
        raise AssertionError("손상된 이미지 파일을 거절해야 합니다.")


def test_cp949_text_is_supported():
    result = extract_document("memo.txt", "한글 메모입니다.".encode("cp949"))
    assert result.text == "한글 메모입니다."
    assert result.warnings == ["텍스트 인코딩을 cp949로 감지했습니다."]


def test_old_doc_format_has_clear_error():
    try:
        extract_document("legacy.doc", b"old word document")
    except UnsupportedDocumentError as exc:
        assert ".docx" in str(exc)
    else:
        raise AssertionError(".doc 파일은 명확한 오류를 반환해야 합니다.")


def test_file_size_limit_is_enforced():
    try:
        extract_document(
            "large.txt",
            b"12345",
            config=ExtractionConfig(max_file_bytes=4),
        )
    except InvalidDocumentError as exc:
        assert "파일이 너무 큽니다" in str(exc)
    else:
        raise AssertionError("크기 제한보다 큰 파일을 거절해야 합니다.")


if __name__ == "__main__":
    tests = [value for name, value in sorted(globals().items()) if name.startswith("test_")]
    passed = 0
    for test in tests:
        try:
            test()
            print(f"  ✅ {test.__name__}")
            passed += 1
        except Exception as error:
            print(f"  ❌ {test.__name__}: {error}")
    print(f"\n{passed}/{len(tests)} 통과")
    if passed != len(tests):
        raise SystemExit(1)
